import os
from typing import List
from langchain.schema import Document


def load_pdf(file_path: str) -> List[Document]:
    from pypdf import PdfReader
    documents = []
    filename = os.path.basename(file_path)
    try:
        reader = PdfReader(file_path)
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                documents.append(Document(
                    page_content=text.strip(),
                    metadata={"source": filename, "page": page_num, "format": "pdf"}
                ))
    except Exception as e:
        raise ValueError(f"Failed to parse PDF '{filename}': {e}")
    return documents


def load_docx(file_path: str) -> List[Document]:
    from docx import Document as DocxDocument
    filename = os.path.basename(file_path)
    doc = DocxDocument(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    full_text = "\n".join(text_parts)
    if full_text.strip():
        return [Document(
            page_content=full_text.strip(),
            metadata={"source": filename, "page": 1, "format": "docx"}
        )]
    return []


def load_txt(file_path: str) -> List[Document]:
    filename = os.path.basename(file_path)
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
    except Exception:
        with open(file_path, "rb") as f:
            text = f.read().decode("utf-8", errors="replace")
    text = text.strip()
    if not text:
        return []
    if len(text) > 5000:
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        return [
            Document(
                page_content=para,
                metadata={"source": filename, "page": i, "format": "txt"}
            )
            for i, para in enumerate(paragraphs, start=1)
        ]
    return [Document(
        page_content=text,
        metadata={"source": filename, "page": 1, "format": "txt"}
    )]


SUPPORTED_EXTENSIONS = {".pdf": load_pdf, ".docx": load_docx, ".txt": load_txt}


def load_document(file_path: str) -> List[Document]:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    loader_fn = SUPPORTED_EXTENSIONS.get(ext)
    if loader_fn is None:
        raise ValueError(f"Unsupported format '{ext}'. Allowed: {', '.join(SUPPORTED_EXTENSIONS)}")
    docs = loader_fn(file_path)
    if not docs:
        raise ValueError(f"No text extracted from '{os.path.basename(file_path)}'.")
    return docs
